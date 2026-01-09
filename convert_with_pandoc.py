# -*- coding: utf-8 -*-
"""
使用 Pandoc 将 Markdown 转换为 Word 文档

Pandoc 安装位置: S:/Tools/Miniconda/envs/pandoc/Library/bin/pandoc.exe
运行方式: 可以用任意 Python 环境运行此脚本

功能：
1. 预处理 Markdown：去掉章节分割符 ---
2. 使用自定义参考模板：正文宋体，代码 Times New Roman
"""

import subprocess
import os
import sys
import tempfile
import re

# Pandoc 可执行文件路径
PANDOC_PATH = r'S:\Tools\Miniconda\envs\pandoc\Library\bin\pandoc.exe'


def preprocess_markdown(content: str) -> str:
    """
    预处理 Markdown 内容

    1. 去掉水平分割线 ---（独立行的三个或更多连字符）

    Args:
        content: 原始 Markdown 内容

    Returns:
        处理后的 Markdown 内容
    """
    # 匹配独立行的水平线：行首可能有空格，然后是3个或更多的 - 或 * 或 _
    # 去掉这些分割线
    lines = content.split('\n')
    processed_lines = []

    for line in lines:
        stripped = line.strip()
        # 检查是否是水平线（3个或更多的 -、* 或 _，可能有空格分隔）
        if re.match(r'^[-*_]{3,}$', stripped.replace(' ', '')):
            # 跳过水平线
            continue
        processed_lines.append(line)

    return '\n'.join(processed_lines)


def create_reference_docx(output_path: str) -> bool:
    """
    创建自定义 Word 参考模板

    样式设置：
    - 正文使用宋体
    - 代码使用 Times New Roman

    Args:
        output_path: 输出模板文件路径

    Returns:
        是否成功创建
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        print("警告: python-docx 未安装，将尝试使用 pandoc 默认模板")
        print("安装方法: pip install python-docx")
        return False

    # 首先用 pandoc 生成默认模板
    temp_template = output_path + '.temp'
    cmd = [
        PANDOC_PATH,
        '-o', temp_template,
        '--print-default-data-file', 'reference.docx'
    ]

    try:
        # 尝试直接生成一个空文档作为模板基础
        result = subprocess.run(
            [PANDOC_PATH, '--print-default-data-file', 'reference.docx'],
            capture_output=True,
            check=False
        )

        if result.returncode == 0:
            with open(temp_template, 'wb') as f:
                f.write(result.stdout)
        else:
            # 如果无法获取默认模板，创建一个最小的 docx
            doc = Document()
            doc.save(temp_template)

    except Exception as e:
        print(f"生成默认模板失败: {e}")
        # 创建一个最小的 docx
        doc = Document()
        doc.save(temp_template)

    # 打开模板并修改样式
    try:
        doc = Document(temp_template)
        styles = doc.styles

        # 设置正文样式 (Normal) - 宋体
        if 'Normal' in [s.name for s in styles]:
            normal_style = styles['Normal']
            normal_style.font.name = '宋体'
            normal_style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            normal_style.font.size = Pt(12)

        # 设置各级标题样式 - 宋体加粗
        for i in range(1, 10):
            heading_name = f'Heading {i}'
            if heading_name in [s.name for s in styles]:
                heading_style = styles[heading_name]
                heading_style.font.name = '宋体'
                heading_style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                heading_style.font.bold = True

        # 设置代码样式 - Times New Roman
        code_style_names = [
            'Source Code',
            'Verbatim Char',
            'Code',
            'HTML Code',
            'Block Text'
        ]

        for style_name in code_style_names:
            if style_name in [s.name for s in styles]:
                code_style = styles[style_name]
                code_style.font.name = 'Times New Roman'
                code_style.font.size = Pt(10)

        # 尝试创建或修改代码相关样式
        try:
            # 修改内置的代码样式
            for style in styles:
                style_name_lower = style.name.lower() if style.name else ''
                if any(x in style_name_lower for x in ['code', 'verbatim', 'source', 'mono']):
                    style.font.name = 'Times New Roman'
                    style.font.size = Pt(10)
        except Exception:
            pass

        # 保存修改后的模板
        doc.save(output_path)

        # 清理临时文件
        if os.path.exists(temp_template):
            os.remove(temp_template)

        print(f"参考模板已创建: {output_path}")
        return True

    except Exception as e:
        print(f"修改模板样式失败: {e}")
        # 清理临时文件
        if os.path.exists(temp_template):
            os.remove(temp_template)
        return False


def convert_md_to_docx(input_file: str, output_file: str, reference_doc: str = None) -> bool:
    """
    使用 pandoc 将 Markdown 转换为 Word 文档

    Args:
        input_file: 输入的 Markdown 文件路径
        output_file: 输出的 Word 文件路径
        reference_doc: 可选的参考模板文档（用于自定义样式）

    Returns:
        是否成功转换
    """

    # 检查输入文件是否存在
    if not os.path.exists(input_file):
        print(f"错误: 输入文件不存在: {input_file}")
        return False

    # 检查 pandoc 是否存在
    if not os.path.exists(PANDOC_PATH):
        print(f"错误: 找不到 pandoc: {PANDOC_PATH}")
        return False

    # 读取并预处理 Markdown 内容
    print("正在预处理 Markdown 文件...")
    with open(input_file, 'r', encoding='utf-8') as f:
        content = f.read()

    processed_content = preprocess_markdown(content)

    # 统计移除的分割线数量
    original_lines = content.count('\n')
    processed_lines = processed_content.count('\n')
    removed_count = original_lines - processed_lines
    if removed_count > 0:
        print(f"  已移除 {removed_count} 个水平分割线")

    # 创建临时文件存储预处理后的内容
    temp_dir = tempfile.gettempdir()
    temp_md = os.path.join(temp_dir, 'temp_processed.md')

    with open(temp_md, 'w', encoding='utf-8') as f:
        f.write(processed_content)

    # 构建 pandoc 命令
    cmd = [
        PANDOC_PATH,
        temp_md,
        '-o', output_file,
        '--from', 'markdown+tex_math_dollars+raw_tex',  # 支持 LaTeX 数学公式
        '--to', 'docx',
        '--standalone',
        '--toc',  # 生成目录
        '--toc-depth=3',  # 目录深度
        # '--number-sections',  # 不自动编号（Markdown 已有编号）
        '--highlight-style', 'tango',  # 代码高亮样式
        '--wrap=auto',
    ]

    # 如果有参考模板，添加到命令中
    if reference_doc and os.path.exists(reference_doc):
        cmd.extend(['--reference-doc', reference_doc])
        print(f"使用参考模板: {reference_doc}")

    print(f"正在转换: {input_file}")
    print(f"输出文件: {output_file}")
    print(f"执行命令: {' '.join(cmd)}")

    try:
        # 执行 pandoc 命令
        result = subprocess.run(
            cmd,
            capture_output=True,
            text=True,
            encoding='utf-8'
        )

        if result.returncode == 0:
            print(f"\n转换成功!")
            print(f"文件已保存到: {output_file}")

            # 显示文件大小
            if os.path.exists(output_file):
                size = os.path.getsize(output_file)
                print(f"文件大小: {size / 1024:.1f} KB")

            # 清理临时文件
            if os.path.exists(temp_md):
                os.remove(temp_md)

            return True
        else:
            print(f"\n转换失败!")
            print(f"错误信息: {result.stderr}")
            return False

    except FileNotFoundError:
        print("错误: 找不到 pandoc 命令，请确保已安装 pandoc")
        print("安装方法: conda install -c conda-forge pandoc")
        return False
    except Exception as e:
        print(f"错误: {str(e)}")
        return False
    finally:
        # 确保清理临时文件
        if os.path.exists(temp_md):
            try:
                os.remove(temp_md)
            except Exception:
                pass


def apply_chinese_fonts_to_docx(docx_path: str) -> bool:
    """
    对生成的 Word 文档应用中文字体设置

    在 pandoc 转换后，进一步确保字体设置正确：
    - 正文：宋体
    - 代码：Times New Roman

    Args:
        docx_path: Word 文档路径

    Returns:
        是否成功应用
    """
    try:
        from docx import Document
        from docx.shared import Pt
        from docx.oxml.ns import qn
    except ImportError:
        print("警告: python-docx 未安装，无法后处理字体")
        return False

    try:
        doc = Document(docx_path)

        # 遍历所有段落
        for para in doc.paragraphs:
            # 检查段落样式
            style_name = para.style.name.lower() if para.style and para.style.name else ''

            for run in para.runs:
                # 检查是否是代码相关样式
                run_style = run.style.name.lower() if run.style and run.style.name else ''

                if any(x in style_name for x in ['code', 'verbatim', 'source']) or \
                   any(x in run_style for x in ['code', 'verbatim', 'source']):
                    # 代码使用 Times New Roman
                    run.font.name = 'Times New Roman'
                    run._element.rPr.rFonts.set(qn('w:ascii'), 'Times New Roman')
                    run._element.rPr.rFonts.set(qn('w:hAnsi'), 'Times New Roman')
                else:
                    # 正文使用宋体
                    run.font.name = '宋体'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

        # 遍历所有表格
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.font.name = '宋体'
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

        doc.save(docx_path)
        print(f"字体设置已应用到: {docx_path}")
        return True

    except Exception as e:
        print(f"应用字体设置失败: {e}")
        return False


def convert_md_to_docx_simple(input_file: str, output_file: str) -> bool:
    """
    简单版本的转换（不带目录和编号）
    """
    cmd = [
        PANDOC_PATH,
        input_file,
        '-o', output_file,
        '--from', 'markdown+tex_math_dollars+raw_tex',
        '--to', 'docx',
        '--standalone',
    ]

    print(f"正在转换 (简单模式): {input_file}")

    try:
        result = subprocess.run(cmd, capture_output=True, text=True, encoding='utf-8')

        if result.returncode == 0:
            print(f"转换成功! 输出: {output_file}")
            return True
        else:
            print(f"转换失败: {result.stderr}")
            return False
    except Exception as e:
        print(f"错误: {str(e)}")
        return False


if __name__ == '__main__':
    # 设置输入输出文件
    input_md = r'G:\work-windows\matlab-bridge\简支梁有限元分析报告.md'
    output_docx = r'G:\work-windows\matlab-bridge\简支梁有限元分析报告_pandoc.docx'
    reference_docx = r'G:\work-windows\matlab-bridge\reference.docx'

    print("=" * 60)
    print("使用 Pandoc 转换 Markdown 到 Word")
    print("=" * 60)
    print()

    # 步骤 1: 创建参考模板
    print("[步骤 1/3] 创建自定义参考模板...")
    template_created = create_reference_docx(reference_docx)
    print()

    # 步骤 2: 转换文档（包含预处理）
    print("[步骤 2/3] 转换 Markdown 到 Word...")
    if template_created:
        success = convert_md_to_docx(input_md, output_docx, reference_docx)
    else:
        print("参考模板创建失败，使用默认样式转换...")
        success = convert_md_to_docx(input_md, output_docx)
    print()

    # 步骤 3: 后处理字体设置
    if success:
        print("[步骤 3/3] 应用字体设置...")
        apply_chinese_fonts_to_docx(output_docx)

    print()
    print("=" * 60)

    if not success:
        # 如果失败，尝试简单模式
        print("\n尝试简单模式...")
        output_docx_simple = r'G:\work-windows\matlab-bridge\简支梁有限元分析报告_pandoc_simple.docx'
        convert_md_to_docx_simple(input_md, output_docx_simple)

    print("\n完成!")
