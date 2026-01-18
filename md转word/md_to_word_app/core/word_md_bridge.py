# -*- coding: utf-8 -*-
"""
Word <-> Markdown 双向转换模块
完全转换模式：保留格式（粗体、斜体、换行、缩进、列表序号）
"""

import re
from pathlib import Path
from typing import List, Dict, Tuple, Optional

from docx import Document
from docx.shared import Pt, Twips
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH


def get_numbering_format(doc, numId: int, level: int) -> dict:
    """
    从 numbering.xml 获取编号格式信息
    返回 {'format': 'decimal'|'chineseCounting'|'upperLetter'|..., 'text': '%1.'}
    """
    try:
        numbering_part = doc.part.numbering_part
        if numbering_part is None:
            return None

        numbering_elem = numbering_part._element

        # 查找 num 元素
        ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
        num_elem = numbering_elem.find(f".//w:num[@w:numId='{numId}']", ns)
        if num_elem is None:
            return None

        # 获取 abstractNumId
        abstract_num_id_elem = num_elem.find("w:abstractNumId", ns)
        if abstract_num_id_elem is None:
            return None
        abstract_num_id = abstract_num_id_elem.get(qn('w:val'))

        # 查找 abstractNum 元素
        abstract_num = numbering_elem.find(f".//w:abstractNum[@w:abstractNumId='{abstract_num_id}']", ns)
        if abstract_num is None:
            return None

        # 查找对应级别的 lvl 元素
        lvl_elem = abstract_num.find(f".//w:lvl[@w:ilvl='{level}']", ns)
        if lvl_elem is None:
            return None

        # 获取编号格式
        numFmt_elem = lvl_elem.find("w:numFmt", ns)
        num_format = numFmt_elem.get(qn('w:val')) if numFmt_elem is not None else 'decimal'

        # 获取编号文本模板
        lvlText_elem = lvl_elem.find("w:lvlText", ns)
        lvl_text = lvlText_elem.get(qn('w:val')) if lvlText_elem is not None else '%1.'

        return {'format': num_format, 'text': lvl_text}
    except Exception:
        return None


def format_number(num: int, format_type: str) -> str:
    """根据格式类型转换数字"""
    if format_type == 'decimal':
        return str(num)
    elif format_type == 'chineseCounting' or format_type == 'ideographTraditional':
        # 中文数字：一、二、三...
        chinese_nums = ['零', '一', '二', '三', '四', '五', '六', '七', '八', '九', '十',
                       '十一', '十二', '十三', '十四', '十五', '十六', '十七', '十八', '十九', '二十']
        return chinese_nums[num] if num < len(chinese_nums) else str(num)
    elif format_type == 'upperLetter':
        # A, B, C...
        return chr(ord('A') + num - 1) if 1 <= num <= 26 else str(num)
    elif format_type == 'lowerLetter':
        # a, b, c...
        return chr(ord('a') + num - 1) if 1 <= num <= 26 else str(num)
    elif format_type == 'upperRoman':
        # I, II, III...
        romans = ['', 'I', 'II', 'III', 'IV', 'V', 'VI', 'VII', 'VIII', 'IX', 'X']
        return romans[num] if num < len(romans) else str(num)
    elif format_type == 'lowerRoman':
        # i, ii, iii...
        romans = ['', 'i', 'ii', 'iii', 'iv', 'v', 'vi', 'vii', 'viii', 'ix', 'x']
        return romans[num] if num < len(romans) else str(num)
    else:
        return str(num)


# 用于跟踪文档级别的编号计数器
_doc_numbering_counters = {}


def get_cell_grid_position(table, row_idx: int, col_idx: int) -> Tuple[int, int]:
    """
    获取单元格在网格中的真实起始位置
    处理合并单元格的情况
    """
    try:
        row = table.rows[row_idx]
        cell = row.cells[col_idx]
        tc = cell._tc

        # 检查是否是垂直合并的延续单元格
        vMerge = tc.tcPr.vMerge if tc.tcPr else None
        if vMerge is not None:
            # val 为 None 或 "continue" 表示是合并的延续
            # val 为 "restart" 表示是合并的起始
            val = vMerge.get(qn('w:val'))
            if val is None or val == 'continue':
                # 向上查找真实起始行
                for search_row_idx in range(row_idx - 1, -1, -1):
                    search_cell = table.rows[search_row_idx].cells[col_idx]
                    search_tc = search_cell._tc
                    if id(search_tc) == id(tc):
                        # 同一个单元格，继续向上
                        continue
                    search_vMerge = search_tc.tcPr.vMerge if search_tc.tcPr else None
                    if search_vMerge is not None:
                        val = search_vMerge.get(qn('w:val'))
                        if val == 'restart':
                            return search_row_idx, col_idx
                    break

        # 检查水平合并 - 通过检查 gridSpan
        # 当前位置可能是合并区域的非起始位置
        # 需要向左查找是否有包含当前位置的合并单元格
        for search_col_idx in range(col_idx - 1, -1, -1):
            search_cell = row.cells[search_col_idx]
            if id(search_cell._tc) == id(tc):
                # 同一个单元格对象，说明当前位置在合并区域内
                # 继续向左查找起始位置
                continue
            else:
                # 不同的单元格，说明 search_col_idx + 1 是合并区域的起始
                return row_idx, search_col_idx + 1

        return row_idx, 0 if col_idx > 0 and id(row.cells[0]._tc) == id(tc) else row_idx, col_idx
    except Exception:
        return row_idx, col_idx


def extract_cell_content_with_format(cell, doc=None, numbering_counters=None) -> str:
    """
    提取单元格内容，保留格式信息
    - 粗体用 **text**
    - 斜体用 *text*
    - 段落之间用换行分隔
    - 保留段落开头的空格（记录实际缩进值）
    - 提取真实的列表编号文本

    Args:
        cell: 单元格对象
        doc: Document 对象，用于获取编号格式
        numbering_counters: 编号计数器字典，格式为 {(numId, level): count}
    """
    if numbering_counters is None:
        numbering_counters = {}

    paragraphs_text = []

    for para in cell.paragraphs:
        para_content = []

        # 获取段落开头的缩进（首行缩进 + 左缩进）
        indent_chars = 0
        try:
            pf = para.paragraph_format
            # 首行缩进
            if pf.first_line_indent:
                # 转换为字符数（粗略估计：1个中文字符约 21 磅）
                indent_pt = pf.first_line_indent.pt
                if indent_pt > 0:
                    indent_chars = max(1, int(indent_pt / 10.5))
            # 左缩进
            if pf.left_indent:
                left_pt = pf.left_indent.pt
                if left_pt > 0:
                    indent_chars += max(1, int(left_pt / 10.5))
        except Exception:
            pass

        indent = ' ' * indent_chars if indent_chars > 0 else ''

        # 检查是否是列表项，并获取真实的编号文本
        list_text = ''
        try:
            pPr = para._element.pPr
            numPr = pPr.numPr if pPr is not None else None
            if numPr is not None and doc is not None:
                # 获取编号 ID 和级别
                ilvl_elem = numPr.ilvl
                numId_elem = numPr.numId

                level = ilvl_elem.val if ilvl_elem is not None else 0
                numId = numId_elem.val if numId_elem is not None else 0

                # 获取编号格式
                fmt_info = get_numbering_format(doc, numId, level)

                if fmt_info:
                    # 更新计数器
                    counter_key = (numId, level)
                    if counter_key not in numbering_counters:
                        numbering_counters[counter_key] = 0
                    numbering_counters[counter_key] += 1

                    # 重置更高级别的计数器
                    keys_to_remove = [k for k in numbering_counters.keys()
                                     if k[0] == numId and k[1] > level]
                    for k in keys_to_remove:
                        del numbering_counters[k]

                    current_num = numbering_counters[counter_key]

                    # 格式化编号
                    formatted_num = format_number(current_num, fmt_info['format'])

                    # 应用编号文本模板，如 "(%1)" -> "(一)"
                    list_text = fmt_info['text'].replace(f'%{level + 1}', formatted_num)
        except Exception:
            pass

        for run in para.runs:
            text = run.text
            if not text:
                continue

            # 将软换行（Shift+Enter）转换为 <br> 标记，避免变成硬换行
            text = text.replace('\n', '<br>')

            # 检查格式并添加标记
            is_bold = run.font.bold is True
            is_italic = run.font.italic is True

            if is_bold and is_italic:
                text = f'***{text}***'
            elif is_bold:
                text = f'**{text}**'
            elif is_italic:
                text = f'*{text}*'

            para_content.append(text)

        # 合并段落内容，保留原始空格
        full_para = ''.join(para_content)

        # 添加列表编号和缩进
        if list_text:
            # 直接使用真实的编号文本
            paragraphs_text.append(f'{indent}{list_text}{full_para}')
        else:
            paragraphs_text.append(indent + full_para)

    # 用换行符连接段落
    return '\n'.join(paragraphs_text)


def get_merged_cell_info(table, row_idx: int, col_idx: int) -> Tuple[int, int, int, int]:
    """
    获取合并单元格的信息：起始行、起始列、行跨度、列跨度
    返回 (start_row, start_col, row_span, col_span)
    """
    try:
        row = table.rows[row_idx]
        cell = row.cells[col_idx]
        tc = cell._tc

        # 获取列跨度 (gridSpan)
        gridSpan = tc.tcPr.gridSpan if tc.tcPr else None
        col_span = gridSpan.val if gridSpan is not None else 1

        # 计算真实起始列（向左查找）
        start_col = col_idx
        for search_col in range(col_idx - 1, -1, -1):
            if id(row.cells[search_col]._tc) == id(tc):
                start_col = search_col
            else:
                break

        # 计算真实起始行（向上查找）
        start_row = row_idx
        vMerge = tc.tcPr.vMerge if tc.tcPr else None
        if vMerge is not None:
            val = vMerge.get(qn('w:val'))
            if val is None:  # continue
                for search_row in range(row_idx - 1, -1, -1):
                    search_cell = table.rows[search_row].cells[start_col]
                    search_tc = search_cell._tc
                    search_vMerge = search_tc.tcPr.vMerge if search_tc.tcPr else None
                    if search_vMerge is not None:
                        search_val = search_vMerge.get(qn('w:val'))
                        if search_val == 'restart':
                            start_row = search_row
                            break
                        elif search_val is None:
                            continue
                    else:
                        break

        # 计算行跨度（向下查找）
        row_span = 1
        if vMerge is not None:
            for search_row in range(start_row + 1, len(table.rows)):
                search_cell = table.rows[search_row].cells[start_col]
                search_tc = search_cell._tc
                search_vMerge = search_tc.tcPr.vMerge if search_tc.tcPr else None
                if search_vMerge is not None:
                    search_val = search_vMerge.get(qn('w:val'))
                    if search_val is None:  # continue
                        row_span += 1
                    else:
                        break
                else:
                    break

        return start_row, start_col, row_span, col_span
    except Exception:
        return row_idx, col_idx, 1, 1


def is_vmerge_continue(tc) -> bool:
    """检查单元格是否是垂直合并的延续（不是起始）"""
    try:
        if tc.tcPr is not None:
            vMerge = tc.tcPr.find(qn('w:vMerge'))
            if vMerge is not None:
                val = vMerge.get(qn('w:val'))
                # val 为 None 或 "continue" 表示是合并的延续
                # val 为 "restart" 表示是合并的起始
                return val is None or val == 'continue'
    except Exception:
        pass
    return False


def word_to_markdown(doc_path: str, md_path: Optional[str] = None) -> Tuple[bool, str]:
    """
    将 Word 文档的所有表格内容转换为 Markdown 格式
    保留格式：粗体、斜体、换行、缩进、合并单元格位置、真实列表编号

    处理合并单元格的策略：
    1. 水平合并：在同一行内，通过比较相邻单元格的 tc 对象来检测
    2. 垂直合并：通过检查 vMerge 属性来判断是否是延续单元格
    """
    try:
        doc = Document(doc_path)
    except Exception as e:
        return False, f"无法打开 Word 文件: {e}"

    if md_path is None:
        md_path = str(Path(doc_path).with_suffix('.md'))

    lines = []
    lines.append(f"# {Path(doc_path).stem}\n")
    lines.append(f"<!-- source: {doc_path} -->\n")

    # 文档级别的编号计数器
    numbering_counters = {}

    for t_idx, table in enumerate(doc.tables):
        lines.append(f"\n## 表格 {t_idx}\n")

        for r_idx, row in enumerate(table.rows):
            lines.append(f"\n### 第 {r_idx} 行\n")

            # 在同一行内跟踪已处理的单元格（用于检测水平合并）
            processed_in_row = set()

            for c_idx, cell in enumerate(row.cells):
                tc = cell._tc
                tc_id = id(tc)

                # 检查水平合并：同一行内是否已经处理过这个 tc
                if tc_id in processed_in_row:
                    continue
                processed_in_row.add(tc_id)

                # 检查垂直合并：是否是延续单元格
                if is_vmerge_continue(tc):
                    # 是垂直合并的延续，跳过
                    continue

                # 提取带格式的内容（传入 doc 和计数器）
                content = extract_cell_content_with_format(cell, doc, numbering_counters)

                # 记录单元格位置和内容
                lines.append(f"<!-- cell:{t_idx},{r_idx},{c_idx} -->\n")
                if content:
                    lines.append(f"{content}\n")
                lines.append(f"<!-- /cell -->\n")

    result = ''.join(lines)

    try:
        with open(md_path, 'w', encoding='utf-8') as f:
            f.write(result)
        return True, md_path
    except Exception as e:
        return False, f"无法写入 Markdown 文件: {e}"


def parse_markdown_cells(content: str) -> List[Dict]:
    """解析 Markdown 中的所有单元格"""
    cells = []

    pattern = r'<!-- cell:(\d+),(\d+),(\d+) -->\n(.*?)<!-- /cell -->'

    for match in re.finditer(pattern, content, re.DOTALL):
        # 不要 strip，保留原始格式
        value = match.group(4)
        # 只去掉末尾的单个换行（我们添加的）
        if value.endswith('\n'):
            value = value[:-1]

        cells.append({
            'pos': (int(match.group(1)), int(match.group(2)), int(match.group(3))),
            'value': value
        })

    return cells


def parse_formatted_text(text: str) -> Tuple[str, List[Dict]]:
    """
    解析带格式标记的文本，返回 (缩进前缀, run 列表)
    每个 run 包含 text, bold, italic

    处理：
    - 开头的空格（缩进）
    - 格式标记 **bold**, *italic*, ***bold+italic***
    """
    # 提取开头的空格作为缩进
    indent_match = re.match(r'^( +)', text)
    indent = ''
    if indent_match:
        indent = indent_match.group(1)
        text = text[len(indent):]

    runs = []

    # 匹配 ***bold+italic***, **bold**, *italic*, 或普通文本
    pattern = r'(\*\*\*(.+?)\*\*\*|\*\*(.+?)\*\*|\*(.+?)\*|([^*]+))'

    for match in re.finditer(pattern, text):
        if match.group(2):  # ***bold+italic***
            runs.append({'text': match.group(2), 'bold': True, 'italic': True})
        elif match.group(3):  # **bold**
            runs.append({'text': match.group(3), 'bold': True, 'italic': False})
        elif match.group(4):  # *italic*
            runs.append({'text': match.group(4), 'bold': False, 'italic': True})
        elif match.group(5):  # plain text
            runs.append({'text': match.group(5), 'bold': False, 'italic': False})

    return indent, runs


def markdown_to_word(md_path: str, template_path: str, output_path: str) -> Tuple[bool, str]:
    """
    根据 Markdown 内容生成 Word 文档
    解析格式标记（**bold**, *italic*）并恢复格式
    """
    try:
        with open(md_path, 'r', encoding='utf-8') as f:
            content = f.read()
    except Exception as e:
        return False, f"无法读取 Markdown 文件: {e}"

    try:
        doc = Document(template_path)
    except Exception as e:
        return False, f"无法打开 Word 模板: {e}"

    cells = parse_markdown_cells(content)

    if not cells:
        return False, "未找到可填充的单元格"

    filled_count = 0
    for cell_data in cells:
        t_idx, r_idx, c_idx = cell_data['pos']
        value = cell_data['value']

        if t_idx < len(doc.tables):
            table = doc.tables[t_idx]
            if r_idx < len(table.rows):
                row = table.rows[r_idx]
                if c_idx < len(row.cells):
                    cell = row.cells[c_idx]
                    fill_cell_with_format(cell, value)
                    filled_count += 1

    try:
        doc.save(output_path)
        return True, f"成功填充 {filled_count} 个单元格"
    except Exception as e:
        return False, f"无法保存 Word 文件: {e}"


def insert_zwsp_for_chinese(text: str) -> str:
    """
    在中文字符之间插入零宽空格，让 Word 可以在中文任意位置换行
    同时保持英文单词完整
    """
    if not text:
        return text

    result = []
    prev_is_cjk = False

    for char in text:
        # 判断是否是 CJK（中日韩）字符
        is_cjk = '\u4e00' <= char <= '\u9fff' or \
                 '\u3400' <= char <= '\u4dbf' or \
                 '\uf900' <= char <= '\ufaff'

        # 如果前一个和当前都是中文，插入零宽空格
        if prev_is_cjk and is_cjk:
            result.append('\u200b')  # 零宽空格

        result.append(char)
        prev_is_cjk = is_cjk

    return ''.join(result)


def clear_paragraph_numbering(para) -> None:
    """清除段落的编号属性，避免重复编号"""
    try:
        pPr = para._element.pPr
        if pPr is not None:
            numPr = pPr.find(qn('w:numPr'))
            if numPr is not None:
                pPr.remove(numPr)
    except Exception:
        pass


def fill_cell_with_format(cell, value: str) -> None:
    """
    填充单元格，解析并恢复格式
    支持 **bold**, *italic*, ***bold+italic***
    支持缩进（编号文本已经是普通文本的一部分）
    """
    if not cell.paragraphs:
        return

    # 获取原有的基础格式（字体名称和大小）
    base_format = {
        'font_name': None,
        'font_size': None,
    }

    first_para = cell.paragraphs[0]
    if first_para.runs:
        first_run = first_para.runs[0]
        base_format['font_name'] = first_run.font.name
        base_format['font_size'] = first_run.font.size

    # 获取段落格式
    para_format = {
        'alignment': first_para.alignment,
        'line_spacing': first_para.paragraph_format.line_spacing,
    }

    # 清空所有段落并移除编号属性
    for para in cell.paragraphs:
        para.clear()
        clear_paragraph_numbering(para)

    # 删除多余段落，只保留第一个
    while len(cell.paragraphs) > 1:
        p = cell.paragraphs[-1]._element
        p.getparent().remove(p)

    # 按换行符分割成多个段落
    paragraphs = value.split('\n') if value else ['']

    for p_idx, para_text in enumerate(paragraphs):
        if p_idx == 0:
            para = cell.paragraphs[0]
        else:
            # 添加新段落
            para = cell.add_paragraph()

        # 清除编号属性（确保新段落也没有编号）
        clear_paragraph_numbering(para)

        # 处理对齐方式：居中保持居中，两端对齐改为左对齐
        if para_format['alignment'] == WD_ALIGN_PARAGRAPH.CENTER:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            # 两端对齐(JUSTIFY)或其他 -> 左对齐
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT

        if para_format['line_spacing']:
            para.paragraph_format.line_spacing = para_format['line_spacing']

        # 解析格式标记
        if para_text:
            indent, runs_data = parse_formatted_text(para_text)

            # 处理缩进
            if indent:
                # 将空格数转换为缩进
                indent_chars = len(indent)
                indent_pt = indent_chars * 10.5  # 大约每个空格 10.5 磅
                try:
                    para.paragraph_format.first_line_indent = Pt(indent_pt)
                except Exception:
                    pass

            for run_data in runs_data:
                text = run_data['text']
                # 处理软换行标记 <br>
                if '<br>' in text:
                    parts = text.split('<br>')
                    for i, part in enumerate(parts):
                        if part:  # 非空部分
                            run = para.add_run(insert_zwsp_for_chinese(part))
                            # 应用基础格式
                            if base_format['font_name']:
                                run.font.name = base_format['font_name']
                                try:
                                    run._element.rPr.rFonts.set(qn('w:eastAsia'), base_format['font_name'])
                                except Exception:
                                    pass
                            if base_format['font_size']:
                                run.font.size = base_format['font_size']
                            # 应用粗体/斜体
                            if run_data['bold']:
                                run.font.bold = True
                            if run_data['italic']:
                                run.font.italic = True
                        # 添加软换行（除了最后一个部分）
                        if i < len(parts) - 1:
                            run = para.add_run()
                            run.add_break()
                else:
                    run = para.add_run(insert_zwsp_for_chinese(text))
                    # 应用基础格式
                    if base_format['font_name']:
                        run.font.name = base_format['font_name']
                        try:
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), base_format['font_name'])
                        except Exception:
                            pass
                    if base_format['font_size']:
                        run.font.size = base_format['font_size']
                    # 应用粗体/斜体
                    if run_data['bold']:
                        run.font.bold = True
                    if run_data['italic']:
                        run.font.italic = True


def get_template_source(md_path: str) -> Optional[str]:
    """从 Markdown 文件中提取源 Word 模板路径"""
    try:
        with open(md_path, 'r', encoding='utf-8') as f:
            content = f.read()

        match = re.search(r'<!-- source: (.+?) -->', content)
        if match:
            return match.group(1)
    except:
        pass
    return None
