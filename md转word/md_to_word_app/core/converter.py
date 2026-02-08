# -*- coding: utf-8 -*-
"""
转换服务 - Markdown 到 Word 的完整转换逻辑
"""

import os
import sys
import subprocess
import re
import shutil
from typing import Tuple, Dict, Any, Optional, List


def find_pandoc() -> str:
    """
    自动查找 Pandoc 可执行文件路径

    按以下优先级查找:
    1. PATH 环境变量中的 pandoc
    2. 常见安装位置 (Conda 环境、系统默认路径等)
    3. 硬编码的 fallback 路径

    Returns:
        Pandoc 可执行文件的路径，如果找不到则返回空字符串
    """
    # 根据平台确定可执行文件名
    is_windows = sys.platform == 'win32'
    exe_name = 'pandoc.exe' if is_windows else 'pandoc'

    # 1. 首先在 PATH 中查找
    pandoc_in_path = shutil.which('pandoc')
    if pandoc_in_path and os.path.isfile(pandoc_in_path):
        return pandoc_in_path

    # 2. 检查常见安装位置
    common_paths: List[str] = []

    if is_windows:
        # Windows 常见路径
        conda_base = os.environ.get('CONDA_PREFIX', '')
        user_profile = os.environ.get('USERPROFILE', '')
        program_files = os.environ.get('PROGRAMFILES', r'C:\Program Files')
        program_files_x86 = os.environ.get('PROGRAMFILES(X86)', r'C:\Program Files (x86)')
        local_app_data = os.environ.get('LOCALAPPDATA', '')

        common_paths = [
            # 当前 Conda 环境
            os.path.join(conda_base, 'Library', 'bin', exe_name) if conda_base else '',
            os.path.join(conda_base, 'Scripts', exe_name) if conda_base else '',
            # 用户级别安装
            os.path.join(local_app_data, 'Pandoc', exe_name) if local_app_data else '',
            os.path.join(user_profile, 'AppData', 'Local', 'Pandoc', exe_name) if user_profile else '',
            # 系统级别安装
            os.path.join(program_files, 'Pandoc', exe_name),
            os.path.join(program_files_x86, 'Pandoc', exe_name),
            # Scoop 安装
            os.path.join(user_profile, 'scoop', 'shims', exe_name) if user_profile else '',
            # Chocolatey 安装
            os.path.join(os.environ.get('ChocolateyInstall', r'C:\ProgramData\chocolatey'), 'bin', exe_name),
            # 硬编码 fallback (原始路径)
            r'S:\Tools\Miniconda\envs\pandoc\Library\bin\pandoc.exe',
        ]
    else:
        # Linux/macOS 常见路径
        home = os.path.expanduser('~')
        conda_base = os.environ.get('CONDA_PREFIX', '')

        common_paths = [
            # 当前 Conda 环境
            os.path.join(conda_base, 'bin', exe_name) if conda_base else '',
            # 系统路径
            '/usr/bin/pandoc',
            '/usr/local/bin/pandoc',
            '/opt/homebrew/bin/pandoc',  # macOS Homebrew (Apple Silicon)
            '/home/linuxbrew/.linuxbrew/bin/pandoc',  # Linux Homebrew
            # 用户级别安装
            os.path.join(home, '.local', 'bin', 'pandoc'),
            os.path.join(home, '.cabal', 'bin', 'pandoc'),
        ]

    # 过滤空字符串并检查路径
    for path in common_paths:
        if path and os.path.isfile(path):
            return path

    # 3. 没有找到，返回空字符串
    return ''


def get_default_pandoc_path() -> str:
    """
    获取默认的 Pandoc 路径

    优先使用自动检测，如果检测失败则返回 fallback 路径
    """
    detected = find_pandoc()
    if detected:
        return detected
    # Fallback: 返回一个合理的默认值（用户需要手动配置）
    if sys.platform == 'win32':
        return 'pandoc.exe'
    return 'pandoc'


# 默认 Pandoc 可执行文件路径 (可在设置中修改)
DEFAULT_PANDOC_PATH = get_default_pandoc_path()


def preprocess_markdown(content: str) -> str:
    """
    预处理 Markdown 内容 - 去掉水平分割线
    """
    lines = content.split('\n')
    processed_lines = []

    for line in lines:
        stripped = line.strip()
        if re.match(r'^[-*_]{3,}$', stripped.replace(' ', '')):
            continue
        processed_lines.append(line)

    return '\n'.join(processed_lines)


def add_table_borders(table) -> None:
    """为表格添加完整的边框"""
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls

    tbl = table._element
    tblPr = tbl.tblPr

    tblBorders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '<w:top w:val="single" w:sz="12" w:space="0" w:color="000000"/>'
        '<w:left w:val="single" w:sz="12" w:space="0" w:color="000000"/>'
        '<w:bottom w:val="single" w:sz="12" w:space="0" w:color="000000"/>'
        '<w:right w:val="single" w:sz="12" w:space="0" w:color="000000"/>'
        '<w:insideH w:val="single" w:sz="12" w:space="0" w:color="000000"/>'
        '<w:insideV w:val="single" w:sz="12" w:space="0" w:color="000000"/>'
        '</w:tblBorders>'
    )

    old_borders = tblPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tblBorders')
    if old_borders is not None:
        tblPr.remove(old_borders)

    tblPr.append(tblBorders)


def create_reference_docx(
    output_path: str,
    chinese_font: str = '宋体',
    code_font: str = 'Times New Roman',
    font_size: float = 12,
    line_spacing: float = 1.5,
    pandoc_path: Optional[str] = None
) -> bool:
    """
    创建自定义 Word 参考模板

    Args:
        output_path: 输出模板文件路径
        chinese_font: 中文字体名称
        code_font: 代码字体名称
        font_size: 正文字体大小 (pt)
        line_spacing: 行间距倍数
        pandoc_path: Pandoc 可执行文件路径 (None 时使用自动检测)
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.oxml.ns import qn
    except ImportError:
        return False

    # 使用传入的路径或自动检测
    actual_pandoc_path = pandoc_path or DEFAULT_PANDOC_PATH

    temp_template = output_path + '.temp'

    try:
        result = subprocess.run(
            [actual_pandoc_path, '--print-default-data-file', 'reference.docx'],
            capture_output=True,
            check=False
        )

        if result.returncode == 0:
            with open(temp_template, 'wb') as f:
                f.write(result.stdout)
        else:
            doc = Document()
            doc.save(temp_template)

    except Exception:
        doc = Document()
        doc.save(temp_template)

    try:
        doc = Document(temp_template)
        styles = doc.styles

        # 正文样式
        from docx.enum.text import WD_LINE_SPACING
        if 'Normal' in [s.name for s in styles]:
            normal_style = styles['Normal']
            normal_style.font.name = chinese_font
            normal_style._element.rPr.rFonts.set(qn('w:eastAsia'), chinese_font)
            normal_style.font.size = Pt(font_size)
            if normal_style.paragraph_format:
                normal_style.paragraph_format.line_spacing = line_spacing

        # 标题样式
        for i in range(1, 10):
            heading_name = f'Heading {i}'
            if heading_name in [s.name for s in styles]:
                heading_style = styles[heading_name]
                heading_style.font.name = chinese_font
                heading_style._element.rPr.rFonts.set(qn('w:eastAsia'), chinese_font)
                heading_style.font.bold = True
                heading_style.font.italic = False
                heading_style.font.color.rgb = RGBColor(0, 0, 0)

        # 目录样式
        from docx.enum.text import WD_LINE_SPACING
        for i in range(1, 10):
            toc_name = f'TOC {i}'
            if toc_name in [s.name for s in styles]:
                toc_style = styles[toc_name]
                toc_style.font.name = code_font
                toc_style._element.rPr.rFonts.set(qn('w:ascii'), code_font)
                toc_style._element.rPr.rFonts.set(qn('w:hAnsi'), code_font)
                toc_style.font.size = Pt(10.5)
                toc_style.font.color.rgb = RGBColor(0, 0, 0)
                if toc_style.paragraph_format:
                    toc_style.paragraph_format.line_spacing = 1.0
                    toc_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

        # 代码样式
        code_style_names = ['Source Code', 'Verbatim Char', 'Code', 'HTML Code', 'Block Text']
        for style_name in code_style_names:
            if style_name in [s.name for s in styles]:
                code_style = styles[style_name]
                code_style.font.name = code_font
                code_style.font.size = Pt(10)

        # 修改所有代码相关样式
        try:
            for style in styles:
                style_name_lower = style.name.lower() if style.name else ''
                if any(x in style_name_lower for x in ['code', 'verbatim', 'source', 'mono']):
                    style.font.name = code_font
                    style.font.size = Pt(10)
        except Exception:
            pass

        doc.save(output_path)

        if os.path.exists(temp_template):
            os.remove(temp_template)

        return True

    except Exception:
        if os.path.exists(temp_template):
            os.remove(temp_template)
        return False


def apply_fonts_to_docx(
    docx_path: str,
    chinese_font: str = '宋体',
    code_font: str = 'Times New Roman',
    font_size: float = 12,
    line_spacing: float = 1.5
) -> bool:
    """
    对生成的 Word 文档应用字体设置和表格边框

    Args:
        docx_path: Word 文档路径
        chinese_font: 中文字体名称
        code_font: 代码字体名称
        font_size: 正文字体大小 (pt)
        line_spacing: 行间距倍数
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.oxml.ns import qn
        from docx.enum.text import WD_LINE_SPACING
    except ImportError:
        return False

    try:
        doc = Document(docx_path)

        for para in doc.paragraphs:
            style_name = para.style.name.lower() if para.style and para.style.name else ''
            is_heading = 'heading' in style_name
            is_toc = 'toc' in style_name
            is_code = any(x in style_name for x in ['code', 'verbatim', 'source'])

            # 设置行间距
            if is_toc:
                para.paragraph_format.line_spacing = 1.0
                para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            elif not is_heading and not is_code:
                para.paragraph_format.line_spacing = line_spacing

            for run in para.runs:
                run_style = run.style.name.lower() if run.style and run.style.name else ''

                if is_code or any(x in run_style for x in ['code', 'verbatim', 'source']):
                    run.font.name = code_font
                    run._element.rPr.rFonts.set(qn('w:ascii'), code_font)
                    run._element.rPr.rFonts.set(qn('w:hAnsi'), code_font)
                elif is_toc:
                    run.font.name = code_font
                    run._element.rPr.rFonts.set(qn('w:ascii'), code_font)
                    run._element.rPr.rFonts.set(qn('w:hAnsi'), code_font)
                    run.font.size = Pt(10.5)
                    run.font.color.rgb = RGBColor(0, 0, 0)
                elif is_heading:
                    run.font.name = chinese_font
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), chinese_font)
                    run.font.color.rgb = RGBColor(0, 0, 0)
                    run.font.italic = False
                    run.font.bold = True
                else:
                    run.font.name = chinese_font
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), chinese_font)
                    run.font.size = Pt(font_size)

        # 表格处理
        for table in doc.tables:
            add_table_borders(table)
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.font.name = chinese_font
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), chinese_font)
                            run.font.size = Pt(font_size)

        doc.save(docx_path)
        return True

    except Exception:
        return False


class ConverterService:
    """Markdown 转 Word 转换服务"""

    def __init__(self):
        from .settings_manager import get_settings_manager
        self._settings = get_settings_manager()
        self.pandoc_path = self._settings.pandoc_path

    def check_pandoc(self) -> Tuple[bool, str]:
        """检查 Pandoc 是否可用"""
        if os.path.exists(self.pandoc_path):
            return True, f"Pandoc 路径: {self.pandoc_path}"
        return False, f"找不到 Pandoc: {self.pandoc_path}"

    def convert(
        self,
        input_file: str,
        output_file: str,
        options: Optional[Dict[str, Any]] = None
    ) -> Tuple[bool, str]:
        """
        转换 Markdown 文件到 Word 文档

        Args:
            input_file: 输入的 Markdown 文件路径
            output_file: 输出的 Word 文件路径
            options: 转换选项
                - generate_toc: 是否生成目录
                - toc_depth: 目录深度
                - highlight_style: 代码高亮样式
                - chinese_font: 中文字体
                - code_font: 代码字体
                - font_size: 正文字体大小 (pt)
                - line_spacing: 行间距倍数
        """
        if options is None:
            options = {}

        generate_toc = options.get('generate_toc', True)
        toc_depth = options.get('toc_depth', 3)
        highlight_style = options.get('highlight_style', 'tango')
        chinese_font = options.get('chinese_font', '宋体')
        code_font = options.get('code_font', 'Times New Roman')
        font_size = options.get('font_size', 12)
        line_spacing = options.get('line_spacing', 1.5)

        # 检查输入文件
        if not os.path.exists(input_file):
            return False, f"输入文件不存在: {input_file}"

        # 检查 Pandoc
        pandoc_ok, pandoc_msg = self.check_pandoc()
        if not pandoc_ok:
            return False, pandoc_msg

        try:
            # 参考模板
            output_dir = os.path.dirname(output_file) or '.'
            reference_docx = os.path.join(output_dir, 'reference_template.docx')

            # 创建参考模板
            template_created = create_reference_docx(
                reference_docx,
                chinese_font=chinese_font,
                code_font=code_font,
                font_size=font_size,
                line_spacing=line_spacing,
                pandoc_path=self.pandoc_path
            )

            # 预处理 Markdown
            with open(input_file, 'r', encoding='utf-8') as f:
                content = f.read()

            processed_content = preprocess_markdown(content)

            input_dir = os.path.dirname(os.path.abspath(input_file))
            temp_md = os.path.join(input_dir, '_temp_processed.md')

            with open(temp_md, 'w', encoding='utf-8') as f:
                f.write(processed_content)

            # 构建 Pandoc 命令
            cmd = [
                self.pandoc_path,
                temp_md,
                '-o', output_file,
                '--from', 'markdown+tex_math_dollars+raw_tex',
                '--to', 'docx',
                '--standalone',
                '--wrap=auto',
                '--resource-path', input_dir,
                '--highlight-style', highlight_style,
            ]

            if generate_toc:
                cmd.append('--toc')
                cmd.append(f'--toc-depth={toc_depth}')

            if template_created and os.path.exists(reference_docx):
                cmd.extend(['--reference-doc', reference_docx])

            # 执行转换
            result = subprocess.run(
                cmd,
                capture_output=True,
                text=True,
                encoding='utf-8',
                cwd=input_dir
            )

            success = result.returncode == 0

            if success:
                # 后处理字体
                apply_fonts_to_docx(
                    output_file,
                    chinese_font=chinese_font,
                    code_font=code_font,
                    font_size=font_size,
                    line_spacing=line_spacing
                )

                file_size = os.path.getsize(output_file)
                size_kb = file_size / 1024
                return True, f"转换成功! 文件大小: {size_kb:.1f} KB"
            else:
                return False, f"Pandoc 转换失败: {result.stderr}"

        except Exception as e:
            return False, f"转换过程中发生错误: {str(e)}"

        finally:
            # 清理临时文件
            if 'temp_md' in locals() and os.path.exists(temp_md):
                try:
                    os.remove(temp_md)
                except Exception:
                    pass

            if 'reference_docx' in locals() and os.path.exists(reference_docx):
                try:
                    os.remove(reference_docx)
                except Exception:
                    pass
